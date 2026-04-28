from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import uuid


def create_resume_docx(data, out_dir="uploads"):

    template = data.get("template", "ats")  # ✅ synced with frontend

    doc = Document()

    # ---------- GLOBAL STYLE ----------
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ================================
    # TEMPLATE HEADER
    # ================================

    if template == "modern":
        add_header_center(doc, data)

    elif template == "creative":
        add_header_left(doc, data)

    else:  # ATS DEFAULT
        add_header_ats(doc, data)

    # ---------- SECTIONS ----------
    add_summary(doc, data)
    add_education(doc, data)
    add_skills(doc, data)
    add_experience(doc, data)
    add_projects(doc, data)
    add_certifications(doc, data)
    add_achievements(doc, data)
    add_extracurricular(doc, data)

    # ---------- SAVE ----------
    os.makedirs(out_dir, exist_ok=True)

    filename = f"{uuid.uuid4().hex}_Resume.docx"
    path = os.path.join(out_dir, filename)

    doc.save(path)

    return path


# ================================
# HEADER TYPES
# ================================

def add_header_center(doc, data):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(data.get("name", ""))
    run.bold = True
    run.font.size = Pt(16)

    add_location(doc, data, center=True)
    add_contact(doc, data, center=True)


def add_header_left(doc, data):
    p = doc.add_paragraph()

    run = p.add_run(data.get("name", ""))
    run.bold = True
    run.font.size = Pt(16)

    add_location(doc, data)
    add_contact(doc, data)


def add_header_ats(doc, data):
    # clean ATS style (no fancy alignment)
    p = doc.add_paragraph()

    run = p.add_run(data.get("name", ""))
    run.bold = True
    run.font.size = Pt(14)

    add_contact(doc, data)


def add_location(doc, data, center=False):
    location = ", ".join(filter(None, [
        data.get("city"),
        data.get("state")
    ]))

    if location:
        p = doc.add_paragraph(location)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_contact(doc, data, center=False):
    items = [
        data.get("phone"),
        data.get("email"),
        data.get("linkedin"),
        data.get("github"),
        data.get("portfolio")
    ]

    text = " | ".join([i for i in items if i])

    if text:
        p = doc.add_paragraph(text)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ================================
# SECTION HELPERS
# ================================

def add_heading(doc, title):
    h = doc.add_paragraph()
    run = h.add_run(title)
    run.bold = True
    run.font.size = Pt(13)


# ================================
# SECTIONS
# ================================

def add_summary(doc, data):
    if data.get("summary"):
        add_heading(doc, "Summary")
        doc.add_paragraph(data["summary"].strip())


def add_education(doc, data):
    if data.get("educations"):
        add_heading(doc, "Education")

        for edu in data["educations"]:
            title = f"{edu.get('degree','')} ({edu.get('duration','')})"

            p = doc.add_paragraph()
            p.add_run(title).bold = True

            inst = edu.get("institution", "")
            if inst:
                line = inst
                if edu.get("cgpa"):
                    line += f" | CGPA: {edu['cgpa']}"
                doc.add_paragraph(line)


def add_skills(doc, data):
    if data.get("skills") or data.get("tools"):
        add_heading(doc, "Technical Skills")

        if data.get("skills"):
            doc.add_paragraph(f"Skills: {data['skills']}")

        if data.get("tools"):
            doc.add_paragraph(f"Tools: {data['tools']}")


def add_experience(doc, data):
    if data.get("experiences"):
        add_heading(doc, "Experience")

        for exp in data["experiences"]:
            title = f"{exp.get('title','')} - {exp.get('company','')} ({exp.get('duration','')})"

            p = doc.add_paragraph()
            p.add_run(title).bold = True

            for line in exp.get("description", "").split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip(), style="List Bullet")


def add_projects(doc, data):
    if data.get("projects"):
        add_heading(doc, "Projects")

        for proj in data["projects"]:
            title = f"{proj.get('title','')} ({proj.get('duration','')})"

            p = doc.add_paragraph()
            p.add_run(title).bold = True

            for line in proj.get("description", "").split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip(), style="List Bullet")


def add_certifications(doc, data):
    if data.get("certificates"):
        add_heading(doc, "Certifications")

        for cert in data["certificates"]:
            text = cert.get("title", "")

            if cert.get("issuer"):
                text += f" - {cert['issuer']}"

            if cert.get("duration"):
                text += f" ({cert['duration']})"

            doc.add_paragraph(text)


def add_achievements(doc, data):
    if data.get("achievements"):
        add_heading(doc, "Achievements")

        content = data["achievements"]

        if isinstance(content, list):
            lines = content
        else:
            lines = content.split("\n")

        for line in lines:
            if line.strip():
                doc.add_paragraph(line.strip(), style="List Bullet")


def add_extracurricular(doc, data):
    extra = data.get("extracurricular_content") or data.get("extracurricular")

    if extra:
        add_heading(doc, "Extracurricular Activities")

        for line in extra.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip(), style="List Bullet")